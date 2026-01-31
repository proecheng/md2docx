# -*- coding: utf-8 -*-
"""
MD2DOCX - Markdown to Word Converter with Formula Support
将Markdown文档转换为Word文档，公式自动转换为公式编辑器格式
支持: 拖放文件到exe / 命令行参数 / 双击后选择文件
"""
import re
import sys
import os
from pathlib import Path

# GUI for file selection
import tkinter as tk
from tkinter import filedialog, messagebox

# Document processing
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter

# ===== OMML Namespace =====
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'m': OMML_NS, 'w': WORD_NS}

def M(tag):
    return f'{{{OMML_NS}}}{tag}'

# ===== MathML to OMML Converter =====
def mathml_to_omml(mathml_element):
    """将MathML元素转换为OMML元素"""
    oMathPara = etree.Element(M('oMathPara'), nsmap=NSMAP)
    oMath = etree.SubElement(oMathPara, M('oMath'))
    for child in mathml_element:
        converted = convert_element(child)
        if converted is not None:
            if isinstance(converted, list):
                for item in converted:
                    oMath.append(item)
            else:
                oMath.append(converted)
    return oMathPara

def convert_element(elem):
    """根据MathML元素类型进行转换"""
    tag = etree.QName(elem.tag).localname if elem.tag else None
    if tag is None:
        return None

    converters = {
        'mrow': convert_mrow, 'mi': convert_mi, 'mn': convert_mn,
        'mo': convert_mo, 'mtext': convert_mtext, 'mfrac': convert_mfrac,
        'msqrt': convert_msqrt, 'mroot': convert_mroot, 'msup': convert_msup,
        'msub': convert_msub, 'msubsup': convert_msubsup, 'munder': convert_munder,
        'mover': convert_mover, 'munderover': convert_munderover,
        'mfenced': convert_mfenced, 'mtable': convert_mtable,
        'mspace': convert_mspace, 'mstyle': convert_mstyle,
        'mpadded': convert_mpadded, 'menclose': convert_menclose,
    }

    converter = converters.get(tag)
    if converter:
        return converter(elem)
    else:
        results = []
        for child in elem:
            converted = convert_element(child)
            if converted is not None:
                if isinstance(converted, list):
                    results.extend(converted)
                else:
                    results.append(converted)
        return results if results else None

def convert_mrow(elem):
    results = []
    for child in elem:
        converted = convert_element(child)
        if converted is not None:
            if isinstance(converted, list):
                results.extend(converted)
            else:
                results.append(converted)
    return results

def create_run(text, italic=True):
    r = etree.Element(M('r'))
    if not italic:
        rPr = etree.SubElement(r, M('rPr'))
        sty = etree.SubElement(rPr, M('sty'))
        sty.set(M('val'), 'p')
    t = etree.SubElement(r, M('t'))
    t.text = text if text else ''
    return r

def convert_mi(elem):
    text = elem.text or ''
    mathvariant = elem.get('mathvariant', '')
    is_function = len(text) > 1 or mathvariant == 'normal'
    return create_run(text, italic=not is_function)

def convert_mn(elem):
    return create_run(elem.text or '', italic=False)

def convert_mo(elem):
    text = elem.text or ''
    char_map = {
        '∑': '∑', '∏': '∏', '∫': '∫', '→': '→', '←': '←', '⇒': '⇒',
        '≤': '≤', '≥': '≥', '≠': '≠', '∈': '∈', '∉': '∉',
        '⋅': '·', '×': '×', '÷': '÷', '∀': '∀', '∃': '∃', '∞': '∞', '∂': '∂',
    }
    return create_run(char_map.get(text, text), italic=False)

def convert_mtext(elem):
    return create_run(elem.text or '', italic=False)

def append_content(parent, content):
    if content is None:
        return
    if isinstance(content, list):
        for item in content:
            if item is not None:
                parent.append(item)
    else:
        parent.append(content)

def convert_mfrac(elem):
    children = list(elem)
    if len(children) < 2:
        return None
    f = etree.Element(M('f'))
    fPr = etree.SubElement(f, M('fPr'))
    typ = etree.SubElement(fPr, M('type'))
    typ.set(M('val'), 'bar')
    num = etree.SubElement(f, M('num'))
    append_content(num, convert_element(children[0]))
    den = etree.SubElement(f, M('den'))
    append_content(den, convert_element(children[1]))
    return f

def convert_msqrt(elem):
    rad = etree.Element(M('rad'))
    radPr = etree.SubElement(rad, M('radPr'))
    degHide = etree.SubElement(radPr, M('degHide'))
    degHide.set(M('val'), '1')
    etree.SubElement(rad, M('deg'))
    e = etree.SubElement(rad, M('e'))
    for child in elem:
        append_content(e, convert_element(child))
    return rad

def convert_mroot(elem):
    children = list(elem)
    if len(children) < 2:
        return convert_msqrt(elem)
    rad = etree.Element(M('rad'))
    etree.SubElement(rad, M('radPr'))
    deg = etree.SubElement(rad, M('deg'))
    append_content(deg, convert_element(children[1]))
    e = etree.SubElement(rad, M('e'))
    append_content(e, convert_element(children[0]))
    return rad

def convert_msup(elem):
    children = list(elem)
    if len(children) < 2:
        return None
    sSup = etree.Element(M('sSup'))
    etree.SubElement(sSup, M('sSupPr'))
    e = etree.SubElement(sSup, M('e'))
    append_content(e, convert_element(children[0]))
    sup = etree.SubElement(sSup, M('sup'))
    append_content(sup, convert_element(children[1]))
    return sSup

def convert_msub(elem):
    children = list(elem)
    if len(children) < 2:
        return None
    sSub = etree.Element(M('sSub'))
    etree.SubElement(sSub, M('sSubPr'))
    e = etree.SubElement(sSub, M('e'))
    append_content(e, convert_element(children[0]))
    sub = etree.SubElement(sSub, M('sub'))
    append_content(sub, convert_element(children[1]))
    return sSub

def convert_msubsup(elem):
    children = list(elem)
    if len(children) < 3:
        return None
    sSubSup = etree.Element(M('sSubSup'))
    etree.SubElement(sSubSup, M('sSubSupPr'))
    e = etree.SubElement(sSubSup, M('e'))
    append_content(e, convert_element(children[0]))
    sub = etree.SubElement(sSubSup, M('sub'))
    append_content(sub, convert_element(children[1]))
    sup = etree.SubElement(sSubSup, M('sup'))
    append_content(sup, convert_element(children[2]))
    return sSubSup

def get_text_content(elem):
    if elem.text:
        return elem.text
    for child in elem:
        text = get_text_content(child)
        if text:
            return text
    return ''

def create_nary(char, sub=None, sup=None, base=None):
    nary = etree.Element(M('nary'))
    naryPr = etree.SubElement(nary, M('naryPr'))
    chr_el = etree.SubElement(naryPr, M('chr'))
    chr_el.set(M('val'), char)
    limLoc = etree.SubElement(naryPr, M('limLoc'))
    limLoc.set(M('val'), 'undOvr')
    sub_el = etree.SubElement(nary, M('sub'))
    if sub is not None:
        append_content(sub_el, convert_element(sub))
    sup_el = etree.SubElement(nary, M('sup'))
    if sup is not None:
        append_content(sup_el, convert_element(sup))
    e = etree.SubElement(nary, M('e'))
    if base is not None:
        append_content(e, convert_element(base))
    return nary

def convert_munder(elem):
    children = list(elem)
    if len(children) < 2:
        return None
    base_text = get_text_content(children[0])
    if base_text in ['∑', '∏', '∫', '⋃', '⋂', 'lim']:
        return create_nary(base_text, sub=children[1])
    limLow = etree.Element(M('limLow'))
    etree.SubElement(limLow, M('limLowPr'))
    e = etree.SubElement(limLow, M('e'))
    append_content(e, convert_element(children[0]))
    lim = etree.SubElement(limLow, M('lim'))
    append_content(lim, convert_element(children[1]))
    return limLow

def create_accent(base_elem, accent_char):
    acc = etree.Element(M('acc'))
    accPr = etree.SubElement(acc, M('accPr'))
    chr_el = etree.SubElement(accPr, M('chr'))
    accent_map = {'^': '̂', '̂': '̂', '~': '̃', '̃': '̃', '¯': '̄', '̄': '̄', '→': '⃗', '⃗': '⃗'}
    chr_el.set(M('val'), accent_map.get(accent_char, '̂'))
    e = etree.SubElement(acc, M('e'))
    append_content(e, convert_element(base_elem))
    return acc

def convert_mover(elem):
    children = list(elem)
    if len(children) < 2:
        return None
    accent_text = get_text_content(children[1])
    if accent_text in ['^', '̂', '~', '̃', '¯', '̄', '→', '⃗']:
        return create_accent(children[0], accent_text)
    limUpp = etree.Element(M('limUpp'))
    etree.SubElement(limUpp, M('limUppPr'))
    e = etree.SubElement(limUpp, M('e'))
    append_content(e, convert_element(children[0]))
    lim = etree.SubElement(limUpp, M('lim'))
    append_content(lim, convert_element(children[1]))
    return limUpp

def convert_munderover(elem):
    children = list(elem)
    if len(children) < 3:
        return None
    base_text = get_text_content(children[0])
    if base_text in ['∑', '∏', '∫', '⋃', '⋂']:
        return create_nary(base_text, sub=children[1], sup=children[2])
    return convert_msubsup(elem)

def convert_mfenced(elem):
    d = etree.Element(M('d'))
    dPr = etree.SubElement(d, M('dPr'))
    begChr = etree.SubElement(dPr, M('begChr'))
    begChr.set(M('val'), elem.get('open', '('))
    endChr = etree.SubElement(dPr, M('endChr'))
    endChr.set(M('val'), elem.get('close', ')'))
    e = etree.SubElement(d, M('e'))
    for child in elem:
        append_content(e, convert_element(child))
    return d

def convert_mtable(elem):
    m_elem = etree.Element(M('m'))
    etree.SubElement(m_elem, M('mPr'))
    for child in elem:
        if etree.QName(child.tag).localname == 'mtr':
            mr = etree.SubElement(m_elem, M('mr'))
            for td in child:
                if etree.QName(td.tag).localname == 'mtd':
                    e = etree.SubElement(mr, M('e'))
                    for td_child in td:
                        append_content(e, convert_element(td_child))
    return m_elem

def convert_mspace(elem):
    return create_run(' ', italic=False)

def convert_mstyle(elem):
    results = []
    for child in elem:
        content = convert_element(child)
        if content is not None:
            if isinstance(content, list):
                results.extend(content)
            else:
                results.append(content)
    return results

def convert_mpadded(elem):
    return convert_mstyle(elem)

def convert_menclose(elem):
    return convert_mstyle(elem)

# ===== LaTeX Processing =====
def preprocess_latex(latex_str):
    """预处理LaTeX"""
    s = latex_str.strip()
    s = re.sub(r'\\text\{([^}]*)\}', r'\\mathrm{\1}', s)
    s = re.sub(r'\\mathbf\{([^}]*)\}', r'\\mathrm{\1}', s)
    s = re.sub(r'\\mathbb\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', s)

    for func in ['softmax', 'clip', 'Attention', 'MultiHead', 'Concat', 'AGG',
                 'ActualSaving', 'ExpectedSaving', 'ComfortViolation', 'SafetyViolation',
                 'CLIP', 'VF', 'ReLU', 'sigmoid', 'tanh', 'argmax', 'argmin']:
        s = s.replace(f'\\{func}', f'\\operatorname{{{func}}}')

    s = re.sub(r'([\u4e00-\u9fff]+)', r'\\text{\1}', s)
    return s

def latex_to_omml(latex_str):
    """LaTeX → OMML"""
    latex_str = latex_str.strip()
    while latex_str.startswith('$'):
        latex_str = latex_str[1:]
    while latex_str.endswith('$'):
        latex_str = latex_str[:-1]
    latex_str = latex_str.strip()

    if not latex_str:
        return None

    latex_str = preprocess_latex(latex_str)

    try:
        mathml_str = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        return mathml_to_omml(mathml_tree)
    except Exception:
        return None

def add_omml_to_paragraph(paragraph, latex_str):
    """将OMML添加到段落"""
    omml = latex_to_omml(latex_str)
    if omml is not None:
        omath_list = omml.findall(f'.//{M("oMath")}')
        if omath_list:
            for omath in omath_list:
                paragraph._element.append(omath)
            return True
    run = paragraph.add_run(latex_str)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(12)
    run.italic = True
    return False

# ===== Math Detection in Text =====
def identify_math_in_text(text):
    """识别文本中的数学表达式"""
    patterns = [
        r'[A-Z][a-zA-Z]*_\{?[a-zA-Z0-9,]+\}?',
        r'[a-zA-Z]_\{?[a-zA-Z0-9,]+\}?\^\{?\([^)]+\)\}?',
        r'[a-zA-Z]\^\{?\([^)]+\)\}?_\{?[a-zA-Z0-9,]+\}?',
        r'[A-Za-zπ][_^]?\{?[\\a-zA-Z0-9]+\}?\([^)]+\)',
        r'[a-zA-Z]_\{?[a-zA-Z0-9]+\}?(?!\()',
        r'p_transferable', r'P_\{?conflict\}?',
        r'[A-Z]\^[A-Z]',
        r'[αβγδεζηθλμνξπρσφχψωΓΔΘΛΞΠΣΦΨΩ][_^]?\{?[a-zA-Z0-9₀-₉]+\}?',
        r'∈\[[^\]]+\]',
        r'(?<![a-zA-Z\u4e00-\u9fff])[αβγδεζηθλμνξπρσφχψωΓΔΘΛΞΠΣΦΨΩε](?![a-zA-Z\u4e00-\u9fff])',
    ]

    combined = '|'.join(f'({p})' for p in patterns)
    segments = [(m.start(), m.end(), m.group()) for m in re.finditer(combined, text)]

    if not segments:
        return []

    segments.sort(key=lambda x: x[0])
    merged = [segments[0]]
    for seg in segments[1:]:
        if seg[0] < merged[-1][1]:
            if seg[1] - seg[0] > merged[-1][1] - merged[-1][0]:
                merged[-1] = seg
        else:
            merged.append(seg)
    return merged

def text_to_latex(text):
    """将识别出的数学文本转为LaTeX"""
    s = text
    greek_map = {
        'α': r'\alpha', 'β': r'\beta', 'γ': r'\gamma', 'δ': r'\delta',
        'ε': r'\epsilon', 'ζ': r'\zeta', 'η': r'\eta', 'θ': r'\theta',
        'λ': r'\lambda', 'μ': r'\mu', 'ν': r'\nu', 'ξ': r'\xi',
        'π': r'\pi', 'ρ': r'\rho', 'σ': r'\sigma', 'φ': r'\phi',
        'χ': r'\chi', 'ψ': r'\psi', 'ω': r'\omega',
        'Γ': r'\Gamma', 'Δ': r'\Delta', 'Θ': r'\Theta', 'Λ': r'\Lambda',
        'Ξ': r'\Xi', 'Π': r'\Pi', 'Σ': r'\Sigma', 'Φ': r'\Phi',
        'Ψ': r'\Psi', 'Ω': r'\Omega',
    }
    for gr, ltx in greek_map.items():
        s = s.replace(gr, ltx)

    subscript_map = {
        '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4',
        '₅': '5', '₆': '6', '₇': '7', '₈': '8', '₉': '9',
        'ₐ': 'a', 'ₑ': 'e', 'ᵢ': 'i', 'ⱼ': 'j', 'ₖ': 'k',
        'ₗ': 'l', 'ₘ': 'm', 'ₙ': 'n', 'ₒ': 'o', 'ₚ': 'p',
        'ᵣ': 'r', 'ₛ': 's', 'ₜ': 't', 'ᵤ': 'u', 'ᵥ': 'v', 'ₓ': 'x'
    }
    for uni, asc in subscript_map.items():
        if uni in s:
            s = s.replace(uni, f'_{asc}')

    s = s.replace('∈', r'\in ')
    return s

# ===== Document Building =====
def set_font(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)

def add_heading(doc, text, level=1):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, '黑体', 22, True)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(24)
    else:
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(level, 12))

def add_formula_paragraph(doc, latex_str):
    """添加独立公式段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    omml = latex_to_omml(latex_str)
    if omml is not None:
        p._element.append(omml)
    else:
        run = p.add_run(latex_str)
        run.font.name = 'Cambria Math'
        run.italic = True
    return p

def process_text_with_math(doc, text, indent=True, bold=False, stats=None):
    """处理包含数学表达式的文本段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)

    dollar_formulas = []
    def replace_dollar(match):
        idx = len(dollar_formulas)
        dollar_formulas.append(match.group(1))
        return f'\x00FORMULA{idx}\x00'

    text_with_placeholders = re.sub(r'\$([^$]+)\$', replace_dollar, text)
    math_segments = identify_math_in_text(text_with_placeholders)

    segments = []
    pos = 0
    for start, end, expr in math_segments:
        if start > pos:
            segments.append((pos, start, 'text', text_with_placeholders[pos:start]))
        segments.append((start, end, 'inline', expr))
        pos = end
    if pos < len(text_with_placeholders):
        segments.append((pos, len(text_with_placeholders), 'text', text_with_placeholders[pos:]))

    if not segments:
        segments = [(0, len(text_with_placeholders), 'text', text_with_placeholders)]

    placeholder_pattern = r'\x00FORMULA(\d+)\x00'

    for _, _, seg_type, content in segments:
        if seg_type == 'text':
            parts = re.split(placeholder_pattern, content)
            i = 0
            while i < len(parts):
                if i % 2 == 0:
                    if parts[i]:
                        run = p.add_run(parts[i])
                        set_font(run, '宋体', 12, bold)
                else:
                    idx = int(parts[i])
                    latex = dollar_formulas[idx]
                    if stats:
                        stats['inline'] += 1
                    add_omml_to_paragraph(p, latex)
                i += 1
        else:
            latex = text_to_latex(content)
            if stats:
                stats['inline'] += 1
            add_omml_to_paragraph(p, latex)
    return p

def add_bullet(doc, text, stats=None):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')

    dollar_formulas = []
    def replace_dollar(match):
        idx = len(dollar_formulas)
        dollar_formulas.append(match.group(1))
        return f'\x00FORMULA{idx}\x00'

    text_clean = re.sub(r'\$([^$]+)\$', replace_dollar, text)
    math_segments = identify_math_in_text(text_clean)

    segments = []
    pos = 0
    for start, end, expr in math_segments:
        if start > pos:
            segments.append(('text', text_clean[pos:start]))
        segments.append(('math', expr))
        pos = end
    if pos < len(text_clean):
        segments.append(('text', text_clean[pos:]))

    if not segments:
        segments = [('text', text_clean)]

    placeholder_pattern = r'\x00FORMULA(\d+)\x00'

    for seg_type, content in segments:
        if seg_type == 'text':
            parts = re.split(placeholder_pattern, content)
            i = 0
            while i < len(parts):
                if i % 2 == 0:
                    if parts[i]:
                        run = p.add_run(parts[i])
                        set_font(run, '宋体', 12)
                else:
                    idx = int(parts[i])
                    latex = dollar_formulas[idx]
                    if stats:
                        stats['inline'] += 1
                    add_omml_to_paragraph(p, latex)
                i += 1
        else:
            latex = text_to_latex(content)
            if stats:
                stats['inline'] += 1
            add_omml_to_paragraph(p, latex)
    return p

# ===== Main Conversion =====
def convert_md_to_docx(input_file, output_file=None):
    """主转换函数"""
    input_path = Path(input_file)
    if not input_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_file}")

    if output_file is None:
        output_file = input_path.with_suffix('.docx')

    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    # 设置样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    lines = md_content.split('\n')
    i = 0
    in_formula = False
    formula_buffer = ''
    in_table = False
    table_rows = []
    stats = {'block': 0, 'inline': 0}

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == '---':
            i += 1
            continue

        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            add_heading(doc, stripped[2:], level=0)
            i += 1
            continue

        if stripped.startswith('## '):
            add_heading(doc, stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith('### '):
            add_heading(doc, stripped[4:], level=2)
            i += 1
            continue

        # 独立公式块
        if stripped.startswith('$$') and not in_formula:
            rest = stripped[2:]
            if rest.endswith('$$') and len(rest) > 2:
                formula = rest[:-2].strip()
                if formula:
                    stats['block'] += 1
                    add_formula_paragraph(doc, formula)
            elif rest.endswith('$$'):
                pass
            else:
                in_formula = True
                formula_buffer = rest
            i += 1
            continue

        if in_formula:
            if stripped.endswith('$$'):
                formula_buffer += ' ' + stripped[:-2]
                formula = formula_buffer.strip()
                if formula:
                    stats['block'] += 1
                    add_formula_paragraph(doc, formula)
                formula_buffer = ''
                in_formula = False
            else:
                formula_buffer += ' ' + stripped
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []

            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue

            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)

            next_is_table = i + 1 < len(lines) and lines[i + 1].strip().startswith('|')
            if not next_is_table and table_rows:
                ncols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = 'Table Grid'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < ncols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            for para in cell.paragraphs:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in para.runs:
                                    set_font(run, '宋体', 10, row_idx == 0)

                in_table = False
                table_rows = []

            i += 1
            continue

        # 项目符号
        if stripped.startswith('- '):
            text = stripped[2:]
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            add_bullet(doc, text, stats)
            i += 1
            continue

        # 普通段落
        text = stripped
        is_bold = text.startswith('**') and '**' in text[2:]
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

        process_text_with_math(doc, text, indent=True, bold=is_bold, stats=stats)
        i += 1

    doc.save(str(output_file))
    return stats, output_file

# ===== GUI =====
def select_file():
    """打开文件选择对话框"""
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title="选择Markdown文件",
        filetypes=[("Markdown文件", "*.md"), ("所有文件", "*.*")]
    )
    root.destroy()
    return file_path

def show_result(success, message):
    """显示结果对话框"""
    root = tk.Tk()
    root.withdraw()
    if success:
        messagebox.showinfo("转换完成", message)
    else:
        messagebox.showerror("转换失败", message)
    root.destroy()

def main():
    """主入口"""
    # 获取输入文件
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        input_file = select_file()
        if not input_file:
            return

    # 检查文件
    if not input_file.lower().endswith('.md'):
        show_result(False, "请选择.md文件!")
        return

    if not os.path.exists(input_file):
        show_result(False, f"文件不存在: {input_file}")
        return

    try:
        stats, output_file = convert_md_to_docx(input_file)
        msg = f"转换成功!\n\n输出文件: {output_file}\n\n独立公式: {stats['block']}\n行内公式: {stats['inline']}\n总计: {stats['block'] + stats['inline']}"
        show_result(True, msg)
    except Exception as e:
        show_result(False, f"转换出错:\n{str(e)}")

if __name__ == '__main__':
    main()
